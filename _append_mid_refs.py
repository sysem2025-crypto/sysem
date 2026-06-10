from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

path = r"H:/Life_OS/02_ProgettiPersonali/06_Sysem/Mod-31 Lista documenti per modulo B  rev 6.docx"
doc = Document(path)

doc.add_page_break()
t = doc.add_paragraph("NORME MID DISPONIBILI IN ARCHIVIO INTERNO")
t.runs[0].bold = True
t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("")
intro = doc.add_paragraph("Riferimenti individuati nella cartella tecnica: H:/Life_OS/01_Lavoro/05_RisorseTecniche/01_Normative/Norme")

items = [
    "Direttiva 2014/32/UE: Direttiva 200_22_CE (MID)/2014_32_UE.pdf",
    "Direttiva MID rifusione: Direttiva 2014-32-UE/direu32_2014.pdf",
    "Recepimento nazionale MID: Decreto Legislativo 19 Maggio 2016 N 84.pdf",
    "Storico MID: Direttiva 200_22_CE (MID)/dirce_22_2004.pdf",
    "Guida software MID: WELMEC 7_2/WELMEC_Guide_7.2_v2021.pdf",
    "Guida software MID (aggiornata): WELMEC 7_2/WELMEC_Guide_7.2_version_v2022.pdf",
    "Guida strumenti automatici: WELMEC 11_1/WELMEC_Guide_11.1_v2020.pdf",
    "Norme settore gas in archivio (se applicabili al prodotto): UNI EN 14236, UNI 12405-3, UNI EN ISO 6976, UNI/TS 11291",
]

for it in items:
    doc.add_paragraph(f"- {it}")

note = doc.add_paragraph("Nota operativa")
note.runs[0].bold = True
doc.add_paragraph("Prima dell'emissione finale della lista documentale, verificare edizione vigente, eventuale sostituzione normativa e applicabilita specifica al tipo strumento (MI-xxx).")

doc.save(path)
print(path)
