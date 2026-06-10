from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

path = r"H:/Life_OS/02_ProgettiPersonali/06_Sysem/Mod-31 Lista documenti per modulo B  rev 6.docx"
doc = Document(path)

def add_bullet(text):
    doc.add_paragraph(f"- {text}")

doc.add_page_break()
t = doc.add_paragraph("ALLEGATO - NORMATIVE MID (RIFERIMENTI E APPLICABILITA')")
t.runs[0].bold = True
t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("")
p = doc.add_paragraph("Scopo")
p.runs[0].bold = True
doc.add_paragraph("Il presente allegato integra la lista documentale del Modulo B con i riferimenti normativi MID applicabili al prodotto, al fine di garantire completezza della pratica di certificazione UE del tipo.")

p = doc.add_paragraph("Riferimenti principali")
p.runs[0].bold = True
for item in [
    "Direttiva 2014/32/UE (MID) e successive modifiche.",
    "Norme armonizzate applicabili alla specifica famiglia di strumenti.",
    "Raccomandazioni e guide WELMEC pertinenti al modulo di valutazione scelto.",
    "Disposizioni nazionali di recepimento vigenti nel Paese di immissione sul mercato.",
]:
    add_bullet(item)

p = doc.add_paragraph("Campo di applicazione")
p.runs[0].bold = True
doc.add_paragraph("Le prescrizioni sono applicate a progettazione, verifiche metrologiche, gestione della produzione e tracciabilita dei componenti che influenzano la conformita MID.")

doc.add_page_break()
t = doc.add_paragraph("DOCUMENTAZIONE RICHIESTA A SUPPORTO DELLA CONFORMITA' MID")
t.runs[0].bold = True
t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("")
for title, body in [
    ("1. Fascicolo tecnico", "Descrizione generale del prodotto, disegni costruttivi, schemi, distinta base e logiche funzionali rilevanti ai fini metrologici."),
    ("2. Analisi dei rischi metrologici", "Valutazione dei rischi che possono influenzare accuratezza, stabilita e ripetibilita della misura."),
    ("3. Piano prove e rapporti di test", "Piani di prova, risultati, criteri di accettazione e tracciabilita agli strumenti campione."),
    ("4. Gestione software rilevante", "Versionamento software/firmware, controllo modifiche e misure di protezione contro alterazioni non autorizzate."),
    ("5. Marcature e informazioni utente", "Bozze etichetta, marcatura CE/MID, identificazione organismo notificato, manuale d'uso e installazione."),
    ("6. Controllo produzione", "Procedure interne di controllo, criteri di collaudo finale e registrazioni di conformita lotto/prodotto."),
]:
    p = doc.add_paragraph(title)
    p.runs[0].bold = True
    doc.add_paragraph(body)

doc.add_page_break()
t = doc.add_paragraph("CHECKLIST OPERATIVA PER RACCOLTA E VALIDAZIONE DOCUMENTI")
t.runs[0].bold = True
t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("")
for item in [
    "Confermare il modulo di valutazione della conformita applicato e relativo perimetro.",
    "Verificare coerenza tra disegni, distinta base e versione firmware oggetto di prova.",
    "Accertare presenza dei rapporti di prova completi e firmati.",
    "Verificare la tracciabilita delle apparecchiature utilizzate per i test metrologici.",
    "Controllare completezza di manuali, istruzioni e indicazioni di marcatura.",
    "Registrare eventuali scostamenti e definire piano azioni correttive con responsabili e date.",
    "Eseguire riesame finale del fascicolo prima della trasmissione all'organismo notificato.",
]:
    add_bullet(item)

p = doc.add_paragraph("Nota")
p.runs[0].bold = True
doc.add_paragraph("I contenuti del presente allegato devono essere adattati al tipo di strumento MID effettivamente certificato e alle eventuali richieste specifiche dell'organismo notificato.")

doc.save(path)
print(path)
